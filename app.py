from flask import Flask, render_template, request, jsonify, redirect, url_for
from flask_cors import CORS
import json
import re
import traceback
import pdfplumber
import docx
import os
import io 
import psycopg2 
from psycopg2.extras import RealDictCursor
from supabase import create_client, Client
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from dotenv import load_dotenv
from datetime import datetime

# Load environment variables
load_dotenv()

# Orchestrator import
try:
    from orchestrator import run_sentinel_analysis 
except ImportError:
    print("❌ CRITICAL: orchestrator.py missing in folder!")

app = Flask(__name__)
CORS(app)
app.secret_key = os.environ.get("SENTINEL_SECRET_KEY", "sentinel-super-secret-2026")

login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = "signin"

print("\n" + "="*60)
print("🚀 RESUME SENTINEL - INITIALIZATION")
print("="*60)

# ==========================================
# 1️⃣ CLOUD DATABASE (SUPABASE) - PRIMARY ⭐
# ==========================================
SUPABASE_URL = os.environ.get("SUPABASE_URL", "https://fdrexxunukdbdyathdiv.supabase.co")
SUPABASE_KEY = os.environ.get("SUPABASE_KEY", "sb_publishable_pM5tfO7H9RsCCAOjNcF-XA_HjJkpZqn")

SUPABASE_CONNECTED = False
supabase = None

try:
    if SUPABASE_URL and SUPABASE_KEY:
        supabase = create_client(SUPABASE_URL, SUPABASE_KEY)
        # Test connection
        response = supabase.table('app_users').select("count", count="exact").limit(1).execute()
        SUPABASE_CONNECTED = True
        print("☁️  [SUPABASE] ✅ Connected successfully!")
    else:
        print("⚠️  [SUPABASE] Credentials not found")
except Exception as e:
    print(f"⚠️  [SUPABASE] Connection failed: {str(e)[:80]}")
    SUPABASE_CONNECTED = False

# ==========================================
# 2️⃣ LOCAL DATABASE (POSTGRESQL) - OPTIONAL
# ==========================================
DB_HOST = os.environ.get("DB_HOST", "localhost")
DB_PORT = os.environ.get("DB_PORT", "5432")
DB_NAME = os.environ.get("DB_NAME", "resume_db_schema")
DB_USER = os.environ.get("DB_USER", "postgres")
DB_PASS = os.environ.get("DB_PASSWORD", "root")

POSTGRESQL_CONNECTED = False

def get_db_connection():
    """Get PostgreSQL connection"""
    try:
        conn = psycopg2.connect(
            host=DB_HOST,
            port=DB_PORT,
            database=DB_NAME,
            user=DB_USER,
            password=DB_PASS
        )
        return conn
    except Exception as e:
        return None

# Test PostgreSQL connection
test_conn = get_db_connection()
if test_conn:
    POSTGRESQL_CONNECTED = True
    print("🐘 [POSTGRESQL] ✅ Connected successfully!")
    test_conn.close()
else:
    print("⚠️  [POSTGRESQL] Not available (optional)")

# ==========================================
# DATABASE SCHEMA INITIALIZATION
# ==========================================
def init_postgresql_database():
    """Initialize PostgreSQL tables (optional)"""
    if not POSTGRESQL_CONNECTED:
        return False
    
    try:
        conn = get_db_connection()
        if conn:
            cursor = conn.cursor()
            
            # Profiles table
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS profiles (
                    id SERIAL PRIMARY KEY,
                    name VARCHAR(255) NOT NULL,
                    resume_text TEXT,
                    jd_text TEXT,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """)
            
            # Scan reports table
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS scan_reports (
                    id SERIAL PRIMARY KEY,
                    candidate_name VARCHAR(255),
                    match_score INTEGER,
                    status VARCHAR(255),
                    ai_analysis JSONB,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """)
            
            # App users table
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS app_users (
                    id SERIAL PRIMARY KEY,
                    full_name VARCHAR(255),
                    username VARCHAR(255) UNIQUE,
                    email VARCHAR(255) UNIQUE NOT NULL,
                    password VARCHAR(255) NOT NULL,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """)
            
            # Insert default admin
            cursor.execute("""
                INSERT INTO app_users (full_name, username, email, password)
                VALUES ('System Admin', 'admin', 'admin@sentinel.ai', 'admin123')
                ON CONFLICT (email) DO NOTHING
            """)
            
            conn.commit()
            cursor.close()
            conn.close()
            print("✅ [POSTGRESQL] Tables initialized!")
            return True
    except Exception as e:
        print(f"⚠️  [POSTGRESQL] Schema init failed: {str(e)[:80]}")
        return False

if POSTGRESQL_CONNECTED:
    init_postgresql_database()

print("="*60 + "\n")

# ==========================================
# AUTHENTICATION MODELS
# ==========================================
# ==========================================
# AUTHENTICATION MODELS
# ==========================================
class User(UserMixin):
    # ✅ FIX: Yahan 'full_name="User"' add kiya hai
    def __init__(self, user_id, email, full_name="User"): 
        self.id = str(user_id)
        self.email = email
        self.full_name = full_name

@login_manager.user_loader
def load_user(user_id):
    """Load user from Supabase (primary) or PostgreSQL (fallback)"""
    # Try Supabase first
    if SUPABASE_CONNECTED:
        try:
            # NAYA: "full_name" ko select mein add kiya
            response = supabase.table('app_users').select("id, email, full_name").eq('id', int(user_id)).execute()
            if response.data and len(response.data) > 0:
                user = response.data[0]
                # NAYA: full_name pass kiya
                return User(user['id'], user['email'], user.get('full_name', 'User'))
        except:
            pass
    
    # Fallback to PostgreSQL
    if POSTGRESQL_CONNECTED:
        conn = get_db_connection()
        if conn:
            try:
                cursor = conn.cursor()
                # NAYA: "full_name" ko SELECT mein add kiya
                cursor.execute("SELECT id, email, full_name FROM app_users WHERE id = %s", (user_id,))
                row = cursor.fetchone()
                cursor.close()
                conn.close()
                # NAYA: row[2] (full_name) pass kiya
                return User(row[0], row[1], row[2]) if row else None
            except:
                return None
    
    return None

@login_manager.unauthorized_handler
def unauthorized():
    """Handle unauthorized access"""
    if request.path.startswith("/api/"):
        return jsonify({"login_required": True}), 401
    return redirect(url_for("signin", next=request.path))

# ==========================================
# 📄 FILE EXTRACTION UTILITIES
# ==========================================
def extract_text_from_pdf(file_stream):
    """Extract text from PDF"""
    try:
        pdf = pdfplumber.open(file_stream)
        text = ""
        for page in pdf.pages:
            extracted = page.extract_text()
            if extracted:
                text += extracted + "\n"
        pdf.close()
        return text.strip()
    except Exception as e:
        print(f"❌ PDF extraction error: {str(e)[:80]}")
        return ""

def extract_text_from_docx(file_stream):
    """Extract text from DOCX with proper error handling"""
    try:
        doc = docx.Document(file_stream)
        text = []
        
        # Extract from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_data.append(cell.text)
                if row_data:
                    text.append(" | ".join(row_data))
        
        return "\n".join(text).strip()
    except Exception as e:
        print(f"❌ DOCX extraction error: {str(e)[:80]}")
        return ""

def extract_text_from_file(file):
    """Universal file extractor"""
    filename = file.filename.lower()
    try:
        file.seek(0)
        content = file.read()
        
        if not content:
            print(f"⚠️ File {filename} is empty")
            return ""
        
        if filename.endswith('.pdf'):
            pdf_stream = io.BytesIO(content)
            return extract_text_from_pdf(pdf_stream)
        
        elif filename.endswith('.docx'):
            file_stream = io.BytesIO(content)
            return extract_text_from_docx(file_stream)
        
        elif filename.endswith('.txt'):
            return content.decode('utf-8').strip()
        
        else:
            print(f"⚠️ Unsupported file type: {filename}")
            return ""
            
    except Exception as e:
        print(f"❌ File extraction error for {filename}: {str(e)[:80]}")
        return ""

# ==========================================
# DATABASE SAVE FUNCTIONS
# ==========================================

def save_report_to_supabase(candidate_name, match_score, status, ai_analysis):
    """Save report to Supabase"""
    if not SUPABASE_CONNECTED:
        return False
    
    try:
        supabase.table('scan_reports').insert({
            "candidate_name": candidate_name,
            "match_score": match_score,
            "status": status,
            "ai_analysis": ai_analysis
        }).execute()
        print("✅ [SUPABASE] Report saved")
        return True
    except Exception as e:
        print(f"⚠️ [SUPABASE] Save failed: {str(e)[:80]}")
        return False

def save_report_to_postgresql(candidate_name, match_score, status, ai_analysis):
    """Save report to PostgreSQL (optional)"""
    if not POSTGRESQL_CONNECTED:
        return False
    
    try:
        conn = get_db_connection()
        if conn:
            cursor = conn.cursor()
            cursor.execute("""
                INSERT INTO scan_reports 
                (candidate_name, match_score, status, ai_analysis)
                VALUES (%s, %s, %s, %s)
            """, (candidate_name, match_score, status, json.dumps(ai_analysis)))
            conn.commit()
            cursor.close()
            conn.close()
            print("✅ [POSTGRESQL] Report saved")
            return True
    except Exception as e:
        print(f"⚠️ [POSTGRESQL] Save failed: {str(e)[:80]}")
        return False

def get_user_from_supabase(email):
    """Get user from Supabase"""
    if not SUPABASE_CONNECTED:
        return None
    
    try:
        response = supabase.table('app_users').select("*").eq('email', email).execute()
        print(f"🔍 Supabase query for {email}: {response.data}")
        if response.data and len(response.data) > 0:
            user = response.data[0]
            print(f"✅ User found: {user}")
            return user
        else:
            print(f"⚠️ No user found for {email}")
    except Exception as e:
        print(f"❌ Supabase query error: {str(e)}")
        pass
    
    return None

def get_user_from_postgresql(email):
    """Get user from PostgreSQL"""
    if not POSTGRESQL_CONNECTED:
        return None
    
    conn = get_db_connection()
    if conn:
        try:
            cursor = conn.cursor(cursor_factory=RealDictCursor)
            cursor.execute("SELECT * FROM app_users WHERE email = %s", (email,))
            user = cursor.fetchone()
            cursor.close()
            conn.close()
            return dict(user) if user else None
        except:
            return None
    
    return None

# ==========================================
# 🚀 CORE ROUTES
# ==========================================

@app.route('/')
def home():
    """Home page"""
    return render_template('index.html', 
                         logged_in=current_user.is_authenticated,
                         supabase_connected=SUPABASE_CONNECTED,
                         postgresql_connected=POSTGRESQL_CONNECTED)

@app.route('/signin', methods=['GET', 'POST'])
def signin():
    """Sign in page"""
    if current_user.is_authenticated:
        return redirect(url_for('home'))

    error = None
    email = request.form.get('email', '').strip() if request.method == 'POST' else ''
    
    if request.method == 'POST':
        password = request.form.get('password', '')
        
        # Try Supabase first
        if SUPABASE_CONNECTED:
            try:
                user = get_user_from_supabase(email)
                if user and user.get('password') == password:
                    login_user(User(user['id'], user['email']))
                    return redirect(request.args.get('next') or url_for('home'))
            except:
                pass
        
        # Fallback to PostgreSQL
        if POSTGRESQL_CONNECTED:
            user = get_user_from_postgresql(email)
            if user and user.get('password') == password:
                login_user(User(user['id'], user['email']))
                return redirect(request.args.get('next') or url_for('home'))
        
        error = "❌ Invalid email or password"

    return render_template('signin.html', error=error, email=email, panel='signin')

@app.route('/register', methods=['POST'])
def register():
    """Register new user"""
    full_name = request.form.get('full_name', '').strip()
    email = request.form.get('email', '').strip()
    username = request.form.get('username', '').strip()
    password = request.form.get('password', '')

    if not all([full_name, email, username, password]):
        return render_template('signin.html',
                             reg_error="All fields are required",
                             panel='register')

    # Try Supabase first
    if SUPABASE_CONNECTED:
        try:
            supabase.table('app_users').insert({
                "full_name": full_name,
                "username": username,
                "email": email,
                "password": password
            }).execute()
            print(f"✅ New user registered in Supabase: {email}")
            return render_template('signin.html',
                                 registered=True,
                                 panel='signin')
        except Exception as e:
            if "duplicate key" in str(e).lower():
                return render_template('signin.html',
                                     reg_error="Email or username already exists",
                                     panel='register')
    
    # Fallback to PostgreSQL
    if POSTGRESQL_CONNECTED:
        conn = get_db_connection()
        if conn:
            try:
                cursor = conn.cursor()
                cursor.execute("""
                    INSERT INTO app_users (full_name, username, email, password)
                    VALUES (%s, %s, %s, %s)
                """, (full_name, username, email, password))
                conn.commit()
                cursor.close()
                conn.close()
                print(f"✅ New user registered in PostgreSQL: {email}")
                return render_template('signin.html',
                                     registered=True,
                                     panel='signin')
            except Exception as e:
                if "duplicate" in str(e).lower():
                    return render_template('signin.html',
                                         reg_error="Email or username already exists",
                                         panel='register')
    
    return render_template('signin.html',
                         reg_error="Registration failed - database unavailable",
                         panel='register')

@app.route('/signout')
def signout():
    """Sign out user"""
    if current_user.is_authenticated:
        logout_user()
    return render_template('sigout.html')

@app.route('/analyze', methods=['POST'])
@login_required
def analyze():
    """Main analysis endpoint"""
    try:
        print("\n📥 Analysis request received...")
        
        # 1. Extract resume text
        resume_text = ""
        if 'resume_file' in request.files and request.files['resume_file'].filename:
            file = request.files['resume_file']
            print(f"📄 Processing file: {file.filename}")
            resume_text = extract_text_from_file(file)
            print(f"✅ Extracted {len(resume_text)} characters from file")
        
        # 2. Fallback to textarea
        if not resume_text or len(resume_text.strip()) < 20:
            resume_text = request.form.get('resume', '').strip()
            if resume_text:
                print(f"✅ Using textarea resume ({len(resume_text)} chars)")
        
        # 3. Get JD
        jd_text = request.form.get('jd', '').strip()
        
        # 4. Validate inputs
        min_length = 50
        if len(resume_text) < min_length:
            raise ValueError(f"Resume too short (min {min_length} chars)")
        if len(jd_text) < min_length:
            raise ValueError(f"Job description too short (min {min_length} chars)")
        
        print(f"📊 Inputs validated: Resume={len(resume_text)}c, JD={len(jd_text)}c")
        
        # 5. Run analysis
        print("🤖 Running AI analysis...")
        raw_output = run_sentinel_analysis(resume_text, jd_text)
        
        # 6. Parse JSON response
        json_match = re.search(r'\{.*\}', str(raw_output), re.DOTALL)
        if not json_match:
            raise ValueError("AI returned invalid format")
        
        result_dict = json.loads(json_match.group())
        
        # 7. Compile final response
        final_response = {
            "score": result_dict.get('score', 0),
            "skills": result_dict.get('skills', []),
            "insight": result_dict.get('insight', "Analysis complete"),
            "questions": result_dict.get('questions', [])
        }
        
        print(f"✅ Analysis complete - Score: {final_response['score']}")
        
        # 8. Save to databases
        candidate_name = request.form.get('candidate_name', 'Unknown Candidate')
        status = "APPROVED" if final_response['score'] >= 80 else "REVIEW_NEEDED"
        
        # Supabase (PRIMARY)
        if SUPABASE_CONNECTED:
            save_report_to_supabase(
                candidate_name,
                final_response['score'],
                status,
                final_response
            )
        
        # PostgreSQL (SECONDARY)
        if POSTGRESQL_CONNECTED:
            save_report_to_postgresql(
                candidate_name,
                final_response['score'],
                status,
                final_response
            )
        
        print("\n✅ Analysis and save complete!\n")
        return jsonify(final_response), 200
        
    except Exception as e:
        error_msg = str(e)
        print(f"\n❌ Analysis error: {error_msg}\n")
        return jsonify({
            "score": 0,
            "skills": [],
            "insight": f"Analysis failed: {error_msg}",
            "questions": []
        }), 400

@app.route('/profiles', methods=['GET', 'POST'])
@login_required
def manage_profiles():
    """Manage saved profiles"""
    
    # Try Supabase first
    if SUPABASE_CONNECTED:
        try:
            if request.method == 'GET':
                response = supabase.table('profiles').select("id, name, created_at").order('created_at', desc=True).execute()
                return jsonify({"profiles": response.data if response.data else []})
            
            if request.method == 'POST':
                data = request.get_json()
                response = supabase.table('profiles').insert({
                    "name": data.get('name'),
                    "resume_text": data.get('resume_text'),
                    "jd_text": data.get('jd_text')
                }).execute()
                
                if response.data:
                    return jsonify({
                        "id": response.data[0]['id'],
                        "message": "✅ Profile saved!"
                    }), 201
        except Exception as e:
            print(f"⚠️ Supabase profile error: {str(e)}")
    
    # Fallback to PostgreSQL
    if POSTGRESQL_CONNECTED:
        conn = get_db_connection()
        if not conn:
            return jsonify({"error": "Cannot connect to database"}), 500
        
        cursor = conn.cursor(cursor_factory=RealDictCursor)
        
        if request.method == 'GET':
            cursor.execute("SELECT id, name, created_at FROM profiles ORDER BY created_at DESC")
            profiles = cursor.fetchall()
            cursor.close()
            conn.close()
            return jsonify({"profiles": [dict(p) for p in profiles]})
        
        if request.method == 'POST':
            try:
                data = request.get_json()
                cursor.execute("""
                    INSERT INTO profiles (name, resume_text, jd_text) 
                    VALUES (%s, %s, %s) 
                    RETURNING id
                """, (data.get('name'), data.get('resume_text'), data.get('jd_text')))
                
                new_id = cursor.fetchone()['id']
                conn.commit()
                cursor.close()
                conn.close()
                
                return jsonify({
                    "id": new_id,
                    "message": "✅ Profile saved!"
                }), 201
            except Exception as e:
                return jsonify({"error": f"Save failed: {str(e)[:50]}"}), 400
    
    return jsonify({"error": "No database available"}), 503

@app.route('/reports')
@login_required
def get_reports():
    """Get all saved reports"""
    try:
        # Try Supabase first
        if SUPABASE_CONNECTED:
            response = supabase.table('scan_reports').select("*").order('created_at', desc=True).limit(100).execute()
            return jsonify({
                "status": "success",
                "count": len(response.data) if response.data else 0,
                "data": response.data if response.data else [],
                "database": "Supabase"
            })
        
        # Fallback to PostgreSQL
        if POSTGRESQL_CONNECTED:
            conn = get_db_connection()
            if conn:
                cursor = conn.cursor(cursor_factory=RealDictCursor)
                cursor.execute("SELECT * FROM scan_reports ORDER BY created_at DESC LIMIT 100")
                reports = cursor.fetchall()
                cursor.close()
                conn.close()
                
                return jsonify({
                    "status": "success",
                    "count": len(reports) if reports else 0,
                    "data": [dict(r) for r in reports] if reports else [],
                    "database": "PostgreSQL"
                })
    except Exception as e:
        return jsonify({
            "status": "error",
            "message": str(e)
        }), 500
    
    return jsonify({
        "status": "error",
        "message": "No database available"
    }), 503

@app.route('/status')
def status():
    """Check database status"""
    return jsonify({
        "supabase": "✅ Connected" if SUPABASE_CONNECTED else "❌ Disconnected",
        "postgresql": "✅ Connected" if POSTGRESQL_CONNECTED else "❌ Disconnected",
        "timestamp": datetime.now().isoformat()
    })

# ==========================================
# ERROR HANDLERS
# ==========================================

@app.errorhandler(404)
def not_found(e):
    return jsonify({"error": "Endpoint not found"}), 404

@app.errorhandler(500)
def server_error(e):
    return jsonify({"error": "Internal server error"}), 500

# ==========================================
# MAIN EXECUTION
# ==========================================

if __name__ == '__main__':
    print("\n" + "="*60)
    print("🚀 RESUME SENTINEL - STARTING SERVER")
    print("="*60)
    print(f"Supabase: {'✅ Connected' if SUPABASE_CONNECTED else '❌ Disconnected'}")
    print(f"PostgreSQL: {'✅ Connected' if POSTGRESQL_CONNECTED else '❌ Disconnected'}")
    print("="*60)
    print("📍 Server running on: http://localhost:5000")
    print("="*60 + "\n")
    
    app.run(debug=True, port=5000, host='0.0.0.0')
